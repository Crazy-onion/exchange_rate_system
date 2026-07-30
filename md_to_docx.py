#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert the exchange-rate system manuals (markdown) to Word .docx.

Handles: H1/H2/H3 headings, **bold** inline, - bullet lists,
> blockquotes, and | markdown tables |. Sets a CJK-friendly font.
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

CJK_FONT = "Microsoft YaHei"


def set_cjk(run):
    run.font.name = CJK_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:ascii"), CJK_FONT)
    rfonts.set(qn("w:hAnsi"), CJK_FONT)


def add_inline(paragraph, text):
    """Render inline **bold** segments into a paragraph."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        set_cjk(run)
        if i % 2 == 1:
            run.bold = True


def shade_cell(cell, color="D9E2F3"):
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def parse_table(rows):
    """rows: list of raw '| a | b |' lines. First=header, second=sep."""
    def split_row(line):
        s = line.strip().strip("|")
        return [c.strip() for c in s.split("|")]

    header = split_row(rows[0])
    data = [split_row(r) for r in rows[2:]]  # skip separator row[1]
    return header, data


def convert(md_path, docx_path):
    doc = Document()

    # Base font for normal + headings
    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)

    heading_sizes = {1: 18, 2: 15, 3: 13}

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # ---- Table block ----
        if is_table_row(line) and i + 1 < n and is_table_row(lines[i + 1]):
            block = []
            while i < n and is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            header, data = parse_table(block)
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except Exception:
                table.style = "Table Grid"
            hdr = table.rows[0].cells
            for j, h in enumerate(header):
                hdr[j].text = ""
                p = hdr[j].paragraphs[0]
                add_inline(p, h)
                for r in p.runs:
                    r.bold = True
                shade_cell(hdr[j])
            for row in data:
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = ""
                    add_inline(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # ---- Heading ----
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            p = doc.add_heading(level=min(level, 3))
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(heading_sizes.get(level, 13))
                set_cjk(run)
            i += 1
            continue

        # ---- Blockquote ----
        if line.startswith(">"):
            text = line.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ---- Bullet list ----
        if line.startswith("- "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
            i += 1
            continue

        # ---- Blank ----
        if line.strip() == "":
            i += 1
            continue

        # ---- Normal paragraph ----
        p = doc.add_paragraph()
        add_inline(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    base = "C:/Users/rfuser/WorkBuddy/2026-07-24-15-41-07/exchange_rate_system"
    convert(f"{base}/用户手册.md", f"{base}/用户手册.docx")
    convert(f"{base}/管理员手册.md", f"{base}/管理员手册.docx")
